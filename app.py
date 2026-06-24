import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import io
import tempfile
import uuid
import requests
import PIL.Image
from streamlit_geolocation import streamlit_geolocation

# ==========================================
# 1. CONFIGURAÇÕES INICIAIS DA PÁGINA
# ==========================================
st.set_page_config(
    page_title="Assistente de Laudos Periciais",
    layout="wide"
)

TIPOS_LAUDO = [
    "Furto",
    "Roubo",
    "Trânsito",
    "Morte suspeita",
    "Homicídio",
    "Identificação veicular",
    "Outro tipo de laudo"
]

TOPICOS_PADRAO = [
    "Histórico",
    "Objetivo",
    "Isolamento e preservação do local",
    "Dos Exames (Do local, Dos veículos, Do cadáver, Dos demais vestígios)",
    "Considerações técnico-científicas",
    "Discussão",
    "Conclusão"
]

# ==========================================
# 2. INICIALIZAÇÃO DE ESTADO (SESSION STATE)
# ==========================================
# Migração do Banco de Frases para a memória da sessão para permitir adições dinâmicas
if "frases_padrao" not in st.session_state:
    st.session_state.frases_padrao = {
        "Local não preservado": "O local de crime não se encontrava idoneamente preservado, havendo sinais claros de alteração na disposição original dos vestígios, o que prejudica parcialmente a análise pericial da dinâmica dos fatos.",
        "Local preservado": "O local encontrava-se adequadamente isolado e preservado pela guarnição da Polícia Militar, garantindo a inalterabilidade dos vestígios até a chegada da equipe pericial.",
        "Constatação de Drogas (Cocaína)": "Para a constatação preliminar da natureza da substância, utilizou-se o teste colorimétrico com reagente de Scott (Tiocianato de Cobalto), o qual apresentou coloração azul intensa, indicativo positivo para a presença de alcaloides (cocaína).",
        "Constatação de Drogas (Maconha)": "Realizou-se o teste colorimétrico com o reagente de Fast Blue B, observando-se a formação de coloração avermelhada/purpúrea, indicativo positivo para a presença de canabinoides (Cannabis sativa L.).",
        "Cadáver - Posição e Decúbito": "O cadáver encontrava-se no solo, em decúbito [dorsal/ventral/lateral], com os membros superiores e inferiores em posição de abandono."
    }


def criar_secao(titulo, nivel=1):
    return {
        "id": str(uuid.uuid4()),
        "titulo": titulo,
        "nivel": nivel,
        "incluir": True,
        "rascunho": "",
        "final": "",
        "fotos": [],
        "fila_audios": []
    }


if "secoes_laudo" not in st.session_state:
    st.session_state.secoes_laudo = [criar_secao(t) for t in TOPICOS_PADRAO]

if "mapa_gps" not in st.session_state:
    st.session_state.mapa_gps = None

for secao in st.session_state.secoes_laudo:
    id_sec = secao["id"]
    if f"txt_rasc_{id_sec}" not in st.session_state:
        st.session_state[f"txt_rasc_{id_sec}"] = secao["rascunho"]
    if f"txt_final_{id_sec}" not in st.session_state:
        st.session_state[f"txt_final_{id_sec}"] = secao["final"]

# ==========================================
# 3. FUNÇÕES AUXILIARES
# ==========================================


@st.cache_data
def carregar_modelos_txt():
    modelo_1, modelo_2 = "Nenhum modelo de estrutura encontrado.", "Nenhum modelo de palavras encontrado."
    if os.path.exists("LAUDO PERICIAL MODELO.txt"):
        with open("LAUDO PERICIAL MODELO.txt", "r", encoding="utf-8") as f:
            modelo_1 = f.read()
    if os.path.exists("MODELO COM PALAVRAS.txt"):
        with open("MODELO COM PALAVRAS.txt", "r", encoding="utf-8") as f:
            modelo_2 = f.read()
    return modelo_1, modelo_2


def transcrever_audio(api_key, audio_file_bytes):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-2.5-flash')

    with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp_audio:
        temp_audio.write(audio_file_bytes)
        temp_path = temp_audio.name

    try:
        arquivo_gemini = genai.upload_file(path=temp_path)
        prompt = "Transcreva este áudio exatamente como foi falado, de forma precisa. Retorne apenas o texto transcrito, sem introduções."
        response = model.generate_content([prompt, arquivo_gemini])
        genai.delete_file(arquivo_gemini.name)
        return response.text.strip()
    except Exception as e:
        st.error(f"Erro na transcrição: {e}")
        return ""
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def processar_texto_ia(api_key, tipo_laudo, topico, rascunho, mod1, mod2):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-2.5-flash')

    prompt = f"""
    Você é um Perito Criminal de campo sênior, responsável pela redação de um laudo oficial. 
    Sua tarefa é analisar o relato (rascunho/transcrição) e reescrevê-lo adotando linguagem estritamente técnica, pericial e objetiva. Não é necessário adicionar o nome do tópico ao texto.

    CONTEXTO:
    - Tipo de Laudo: {tipo_laudo}
    - Seção Atual: {topico}

    REGRAS INEGOCIÁVEIS (SIGA RIGOROSAMENTE):
    1. PROIBIDO INVENTAR DADOS: Você NÃO PODE criar, deduzir, presumir ou inventar nenhuma informação, medida, placa, nome, lesão ou vestígio que não esteja explicitamente mencionado no RASCUNHO. Limite-se aos fatos fornecidos.
    2. PERTINÊNCIA E ESCOPO: Escreva SOMENTE informações relativas ao tipo de laudo solicitado e que se enquadrem especificamente na seção "{topico}".
    3. CORREÇÃO LINGUÍSTICA: Identifique e elimine completamente quaisquer gírias, linguagens coloquiais, cacoetes da fala, opiniões pessoais ou achismos.
    4. ESTRUTURA TÉCNICA: Estruture os parágrafos de forma lógica, fluida e formal, preferencialmente utilizando a voz passiva ou terceira pessoa (ex: "Constatou-se", "Observou-se").
    5. JARGÃO PERICIAL: Você DEVE fundamentar seu vocabulário, termos técnicos e estilo de redação baseando-se EXCLUSIVAMENTE nos modelos fornecidos abaixo.

    MODELO DE ESTRUTURA DE LAUDO PERICIAL:
    {mod1}

    MODELO DE PALAVRAS, JARGÕES E ESTILO:
    {mod2}

    RASCUNHO / RELATO DO PERITO:
    "{rascunho}"

    Com base estritamente no rascunho acima e aplicando as regras estabelecidas, gere o texto final para a seção:
    """
    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        st.error(f"Erro na geração de texto: {e}")
        return ""


def gerar_legenda_foto_ia(api_key, image_bytes):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-2.5-flash')
    img = PIL.Image.open(io.BytesIO(image_bytes))
    prompt = "Você é um perito criminal de campo. Crie uma legenda curta, técnica e puramente descritiva (1 frase) para esta foto de local de crime ou vestígio. Não faça deduções, descreva o que é visível com jargão pericial."
    try:
        response = model.generate_content([prompt, img])
        return response.text.strip()
    except Exception as e:
        return f"Erro ao gerar legenda: {e}"


def buscar_mapa_satelite(lat, lon, api_key_maps=""):
    if api_key_maps:
        url = f"https://maps.googleapis.com/maps/api/staticmap?center={lat},{lon}&zoom=18&size=600x400&maptype=satellite&markers=color:red%7C{lat},{lon}&key={api_key_maps}"
        resp = requests.get(url)
        if resp.status_code == 200:
            return resp.content
    else:
        url = f"https://static-maps.yandex.ru/1.x/?ll={lon},{lat}&z=16&l=sat,skl&size=600,400&pt={lon},{lat},pm2rdl"
        try:
            resp = requests.get(url)
            if resp.status_code == 200:
                return resp.content
        except:
            pass
    return None


def gerar_documento_word(tipo_laudo, arquivo_template=None):
    if arquivo_template is not None:
        doc = Document(io.BytesIO(arquivo_template.getvalue()))
    else:
        doc = Document()
        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = p_titulo.add_run('LAUDO PERICIAL\n\n')
        run_titulo.font.name, run_titulo.font.size, run_titulo.font.bold = 'Arial', Pt(
            12), True

        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run(f'Natureza da Ocorrência: {tipo_laudo}')
        run_sub.font.name, run_sub.font.size, run_sub.font.bold = 'Arial', Pt(
            12), True
        doc.add_paragraph()

    contador_topico = 0
    contador_subtopico = 0
    contador_foto = 1
    mapa_inserido = False

    for secao in st.session_state.secoes_laudo:
        if not secao.get('incluir', True):
            continue

        if "exames" in secao['titulo'].lower() and st.session_state.mapa_gps and not mapa_inserido:
            p_mapa = doc.add_paragraph()
            p_mapa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_mapa = p_mapa.add_run()
            run_mapa.add_picture(io.BytesIO(
                st.session_state.mapa_gps['bytes']), height=Cm(10.5))

            p_leg_mapa = doc.add_paragraph()
            p_leg_mapa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_leg_mapa = p_leg_mapa.add_run(
                f"Imagem {contador_foto} - {st.session_state.mapa_gps['legenda']}")
            run_leg_mapa.font.name, run_leg_mapa.font.size = 'Arial', Pt(10)
            contador_foto += 1
            mapa_inserido = True
            doc.add_paragraph()

        # Fallback de conteúdo: Usa o 'final' se existir, caso contrário usa o 'rascunho'
        texto_final = secao['final'].strip(
        ) if secao['final'].strip() else secao['rascunho'].strip()
        fotos = secao['fotos']

        if texto_final or fotos or secao['titulo']:
            if secao['nivel'] == 1:
                contador_topico += 1
                contador_subtopico = 0
                titulo_exibicao = f"{contador_topico}. {secao['titulo'].upper()}"
                tamanho_fonte = 14
            else:
                contador_subtopico += 1
                titulo_exibicao = f"{contador_topico}.{contador_subtopico}. {secao['titulo']}"
                tamanho_fonte = 13

            p_heading = doc.add_paragraph()
            run_heading = p_heading.add_run(f"{titulo_exibicao}\n")
            run_heading.font.name, run_heading.font.size, run_heading.font.bold = 'Arial', Pt(
                tamanho_fonte), True

            if texto_final:
                for linha in texto_final.split('\n'):
                    if linha.strip():
                        p_text = doc.add_paragraph(
                            linha.strip().replace("  ", " "))
                        p_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        for run in p_text.runs:
                            run.font.name, run.font.size = 'Arial', Pt(12)

            if fotos:
                for foto_dict in fotos:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_after = Pt(0)

                    run_img = p_img.add_run()
                    run_img.add_picture(io.BytesIO(
                        foto_dict['bytes']), width=Cm(14.67))

                    p_legenda = doc.add_paragraph()
                    p_legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_legenda.paragraph_format.space_before = Pt(0)

                    legenda_texto = foto_dict.get('legenda', '').strip()
                    if not legenda_texto:
                        legenda_texto = "_________________________"

                    run_legenda = p_legenda.add_run(
                        f"Imagem {contador_foto} - {legenda_texto}")
                    run_legenda.font.name, run_legenda.font.size = 'Arial', Pt(
                        10)
                    contador_foto += 1

    arquivo_io = io.BytesIO()
    doc.save(arquivo_io)
    arquivo_io.seek(0)
    return arquivo_io


# ==========================================
# 4. INTERFACE PRINCIPAL (UI)
# ==========================================
st.title("Assistente de Laudos")

modelo_1_texto, modelo_2_texto = carregar_modelos_txt()

with st.sidebar:
    st.header("Controle do Laudo")

    with st.expander("Configurações de Integração", expanded=False):
        chave_api = st.text_input(
            "Chave API (Google Gemini):", type="password")
        chave_maps = st.text_input("Chave API (Google Maps - Opcional):",
                                   type="password", help="Garante mapa com satélite de alta qualidade.")

    with st.expander("Georreferenciamento", expanded=False):
        st.markdown(
            "Obtenha as coordenadas geográficas. A imagem de satélite será anexada antes da seção **Dos Exames**.")
        localizacao = streamlit_geolocation()
        if localizacao and localizacao.get('latitude'):
            lat, lon = localizacao['latitude'], localizacao['longitude']
            st.success(f"Latitude: {lat:.5f} | Longitude: {lon:.5f}")

            if st.button("Anexar Mapa ao Documento", use_container_width=True):
                with st.spinner("Realizando busca de satélite..."):
                    img_mapa = buscar_mapa_satelite(lat, lon, chave_maps)
                    if img_mapa:
                        st.session_state.mapa_gps = {
                            "bytes": img_mapa,
                            "legenda": f"Vista de satélite do local georreferenciado. Coordenadas obtidas: {lat}, {lon}."
                        }
                        st.success("Mapa inserido no escopo do laudo.")
                    else:
                        st.error("Falha ao obter imagem de satélite.")

        if st.session_state.mapa_gps:
            st.image(
                st.session_state.mapa_gps['bytes'], caption="Mapa Capturado")
            if st.button("Remover Mapa", use_container_width=True):
                st.session_state.mapa_gps = None
                st.rerun()

    with st.expander("Informações do Laudo", expanded=False):
        tipo_laudo_selecionado = st.selectbox(
            "Natureza da Ocorrência:", TIPOS_LAUDO)
        template_upload = st.file_uploader(
            "Documento Padrão - Base (.docx)", type=["docx"])

    with st.expander("Estrutura do Documento", expanded=False):
        st.markdown("**Adicionar Novo Tópico**")
        novo_topico = st.text_input("Título do Novo Tópico Principal:")

        topicos_principais_atual = [
            s for s in st.session_state.secoes_laudo if s['nivel'] == 1]
        opcoes_posicao = ["Inserir no início"] + \
            [f"Após: {t['titulo']}" for t in topicos_principais_atual]
        posicao_selecionada = st.selectbox(
            "Posição de Inserção:", opcoes_posicao)

        if st.button("Inserir Tópico", use_container_width=True):
            if novo_topico.strip():
                nova_sec = criar_secao(novo_topico.strip(), nivel=1)
                if posicao_selecionada == "Inserir no início":
                    st.session_state.secoes_laudo.insert(0, nova_sec)
                else:
                    titulo_alvo = posicao_selecionada.replace("Após: ", "")
                    idx_insercao = len(st.session_state.secoes_laudo)
                    for i, sec in enumerate(st.session_state.secoes_laudo):
                        if sec['titulo'] == titulo_alvo and sec['nivel'] == 1:
                            j = i + 1
                            while j < len(st.session_state.secoes_laudo) and st.session_state.secoes_laudo[j]['nivel'] > 1:
                                j += 1
                            idx_insercao = j
                            break
                    st.session_state.secoes_laudo.insert(
                        idx_insercao, nova_sec)
                st.rerun()

        st.markdown("---")
        st.markdown("**Adicionar Novo Subtópico**")
        if topicos_principais_atual:
            opcoes_pai = {s['id']: s['titulo']
                          for s in topicos_principais_atual}
            pai_selecionado = st.selectbox("Vincular subtópico à seção:", options=list(
                opcoes_pai.keys()), format_func=lambda x: opcoes_pai[x])
            novo_subtopico = st.text_input("Título do Subtópico:")
            if st.button("Inserir Subtópico", use_container_width=True):
                if novo_subtopico.strip():
                    idx_pai = next(i for i, s in enumerate(
                        st.session_state.secoes_laudo) if s['id'] == pai_selecionado)
                    idx_insercao = idx_pai + 1
                    while idx_insercao < len(st.session_state.secoes_laudo) and st.session_state.secoes_laudo[idx_insercao]['nivel'] > 1:
                        idx_insercao += 1
                    nova_subsec = criar_secao(novo_subtopico.strip(), nivel=2)
                    st.session_state.secoes_laudo.insert(
                        idx_insercao, nova_subsec)
                    st.rerun()

    with st.expander("Textos Padrão (Snippets)", expanded=False):
        st.markdown(
            "Alimente o banco de dados com jargões e textos frequentemente utilizados.")
        novo_titulo_snippet = st.text_input("Nome de Referência:")
        novo_texto_snippet = st.text_area("Conteúdo Técnico:")
        if st.button("Salvar Texto Padrão", use_container_width=True):
            if novo_titulo_snippet.strip() and novo_texto_snippet.strip():
                st.session_state.frases_padrao[novo_titulo_snippet.strip(
                )] = novo_texto_snippet.strip()
                st.success("Texto padrão registrado com sucesso.")
                st.rerun()

    with st.expander("Exportação", expanded=True):
        if st.button("Gerar Documento Word", type="primary", use_container_width=True):
            docx_bytes = gerar_documento_word(
                tipo_laudo_selecionado, template_upload)
            st.download_button(
                label="Baixar Arquivo Final (.docx)", data=docx_bytes,
                file_name=f"laudo_{tipo_laudo_selecionado.replace(' ', '_').lower()}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

if not chave_api:
    st.info("Insira a chave de integração da API do Gemini na barra lateral para ativar funções automáticas.")

st.markdown("### Preenchimento de Seções do Laudo")

for idx_secao, secao in enumerate(st.session_state.secoes_laudo):
    sec_id = secao['id']
    prefixo_exibicao = "Seção:" if secao['nivel'] == 1 else "Subseção:"

    with st.expander(f"{prefixo_exibicao} {secao['titulo']}", expanded=False):
        c_head1, c_head2 = st.columns([4, 1])
        with c_head1:
            secao['incluir'] = st.checkbox("Incluir esta seção no documento final", value=secao.get(
                'incluir', True), key=f"check_{sec_id}")
        with c_head2:
            if st.button("Excluir Seção", key=f"del_sec_{sec_id}"):
                st.session_state.secoes_laudo.pop(idx_secao)
                st.rerun()

        if secao['incluir']:
            col_esq, col_dir = st.columns([1, 1])

            with col_esq:
                st.markdown("**1. Captação de Dados**")

                audio_gravado = st.audio_input(
                    f"Gravação de áudio em campo", key=f"audio_{sec_id}")
                if audio_gravado:
                    if st.button("Adicionar à Fila de Transcrição", key=f"add_fila_{sec_id}"):
                        secao['fila_audios'].append({"id": str(uuid.uuid4(
                        )), "nome": f"Registro de Áudio {len(secao['fila_audios']) + 1}", "bytes": audio_gravado.getvalue()})
                        st.success("Áudio registrado na fila.")
                        st.rerun()

                if secao['fila_audios']:
                    for idx_aud, aud in enumerate(secao['fila_audios']):
                        ca1, ca2 = st.columns([4, 1])
                        ca1.caption(f"Arquivo pendente: {aud['nome']}")
                        if ca2.button("Remover", key=f"del_aud_{sec_id}_{aud['id']}"):
                            secao['fila_audios'].pop(idx_aud)
                            st.rerun()
                    if st.button("Processar Fila de Transcrição", key=f"btn_transc_{sec_id}", type="primary"):
                        if chave_api:
                            with st.spinner("Processando áudios via IA..."):
                                novos = [transcrever_audio(
                                    chave_api, a['bytes']) for a in secao['fila_audios']]
                                validos = [t for t in novos if t]
                                if validos:
                                    combo = f"{st.session_state[f'txt_rasc_{sec_id}']}\n\n" + "\n\n".join(
                                        validos)
                                    st.session_state[f"txt_rasc_{sec_id}"] = combo.strip(
                                    )
                                    secao['fila_audios'].clear()
                                    st.rerun()
                        else:
                            st.error("Chave API não informada.")

                st.markdown("**Inserção Rápida de Textos Padrão:**")
                snip_sel = st.selectbox("Selecione um texto técnico", ["Selecionar texto..."] + list(
                    st.session_state.frases_padrao.keys()), key=f"sel_snip_{sec_id}", label_visibility="collapsed")
                if st.button("Inserir Texto", key=f"btn_snip_{sec_id}"):
                    if snip_sel != "Selecionar texto...":
                        st.session_state[f"txt_rasc_{sec_id}"] += f"\n{st.session_state.frases_padrao[snip_sel]}"
                        st.rerun()

                secao['rascunho'] = st.text_area(
                    "Texto Rascunho / Anotações:", value=st.session_state[f"txt_rasc_{sec_id}"], height=150, key=f"txt_rasc_{sec_id}")

            with col_dir:
                st.markdown("**2. Estruturação Técnica**")
                if st.button("Gerar Texto Técnico (IA)", key=f"btn_ia_{sec_id}", type="secondary"):
                    if not chave_api:
                        st.error("Chave API não informada.")
                    elif not secao['rascunho'].strip():
                        st.warning(
                            "O rascunho está vazio. Insira informações antes de processar.")
                    else:
                        with st.spinner("Formatando estrutura e linguagem técnica..."):
                            txt = processar_texto_ia(
                                chave_api, tipo_laudo_selecionado, secao['titulo'], secao['rascunho'], modelo_1_texto, modelo_2_texto)
                            if txt:
                                st.session_state[f"txt_final_{sec_id}"] = txt
                                st.rerun()

                secao['final'] = st.text_area(
                    "Texto Final Revisado:", value=st.session_state[f"txt_final_{sec_id}"], height=250, key=f"txt_final_{sec_id}")

            st.divider()
            st.markdown("**Galeria de Registros Fotográficos**")
            fotos_up = st.file_uploader("Selecionar Imagens Locais", type=[
                                        "jpg", "jpeg", "png"], accept_multiple_files=True, key=f"up_{sec_id}")
            if st.button("Armazenar Imagens", key=f"btn_salv_{sec_id}"):
                if fotos_up:
                    for f in fotos_up:
                        secao['fotos'].append(
                            {'id': str(uuid.uuid4()), 'bytes': f.getvalue(), 'legenda': ''})
                    st.rerun()

            if secao['fotos']:
                cols = st.columns(3)
                for i_foto, f_dict in enumerate(secao['fotos']):
                    with cols[i_foto % 3]:
                        st.image(f_dict['bytes'], use_container_width=True)
                        nova_leg = st.text_area(
                            "Legenda Técnica", value=f_dict['legenda'], key=f"leg_{sec_id}_{f_dict['id']}", height=80)
                        f_dict['legenda'] = nova_leg

                        cg1, cg2 = st.columns([3, 1])
                        if cg1.button("Sugerir Legenda (IA)", key=f"btn_leg_ia_{sec_id}_{f_dict['id']}"):
                            if chave_api:
                                with st.spinner("Analisando vestígios visuais..."):
                                    f_dict['legenda'] = gerar_legenda_foto_ia(
                                        chave_api, f_dict['bytes'])
                                    st.rerun()
                            else:
                                st.error("Chave API não informada.")
                        if cg2.button("Excluir", key=f"del_f_{sec_id}_{f_dict['id']}", help="Remover imagem do laudo"):
                            secao['fotos'].remove(f_dict)
                            st.rerun()
